from bs4 import BeautifulSoup
from urllib.parse import urljoin
from tabulate import tabulate

from docx import Document
from docx.shared import Inches

import io
import os
import requests
from PIL import Image

doc = Document()
doc.add_heading('Django Web-Scraping Application', level=1)
doc.add_paragraph('This Application developed by @naveenarikela')

SPECIAL_CONTAINER_TAGS = ['body']
SPLIT_LINES = '\n'
FILE_EXTENSION = '.docx'

#Store images & meta in Local diretory
STATIC_FILES_OUTPUT_DIRECTORY = 'static\images'

OUTPUT_FILES_DIRECTORY = 'output_files'

os.makedirs(STATIC_FILES_OUTPUT_DIRECTORY, exist_ok=True)
os.makedirs(OUTPUT_FILES_DIRECTORY, exist_ok=True)

class ScrapWebiteContent:
    def __init__(self, domain_name, container_tag='body', tags=[], output_filename='web_scrape'):
        self.domain_name = domain_name
        self.container_tag = container_tag
        self.tags = tags
        self.soup_object = self.get_soup_object()
        self.output_file = output_filename + FILE_EXTENSION
        
    def get_soup_object(self):
        html_content = requests.get(self.domain_name).text
        soup = BeautifulSoup(html_content, 'lxml')
        return soup
    
    def get_container_tag(self):
        if self.container_tag in SPECIAL_CONTAINER_TAGS:
            #NOTE:: Single Parent tag, all tags are wrapped inside it.
            container_tag = self.soup_object.find(self.container_tag)
            #to hanle iterations
            container_tag = [container_tag]
        else:
            #NOTE:: Child tags contains more than one. Eg: div, ul...
            container_tag = self.soup_object.find_all(self.container_tag)

        return container_tag
    
    def process_paragraph_tags(self, tag):
        processed_paragraph_tags = ''
        paragraph_tags = self.container_tag_data.find_all(tag)
        for paragraph_tag in paragraph_tags:
            # print(f'paragraph_tag:: {paragraph_tag}')
            processed_paragraph_tags += paragraph_tag.text + SPLIT_LINES 

        doc.add_paragraph(processed_paragraph_tags)
        # doc.save(self.output_file)
        return processed_paragraph_tags
    
    #NOTE:: To handle huge requests we need to write this logic in async after sometime
    def process_image_tags(self, tag):
        image_tags = self.container_tag_data.find_all(tag)
        for image_tag in image_tags:
            # print(f"image_tag:: {image_tag}")
            try:
                img_src = image_tag.get('src')
                img_url = urljoin(self.domain_name, img_src)
                img_data = requests.get(img_url).content
                img_filename = os.path.join(STATIC_FILES_OUTPUT_DIRECTORY, os.path.basename(img_url))
                # print(f"img_filename: {img_filename}")

                #Save images in local directory
                with open(img_filename, 'wb') as img_file:
                    img_file.write(img_data)
                    # print(f"Image saved in local directory")

                #Open and write to docx in png format
                with Image.open(img_filename) as img_file:
                    img_byte_array = io.BytesIO()
                    img_file.save(img_byte_array, format='PNG')
                    doc.add_picture(img_byte_array, width=Inches(2))

                doc.add_paragraph(SPLIT_LINES)
            except Exception as e:
                print(f'WARNING!!:: Unsupported image format Exception: {e}')
                continue

            # doc.save(self.output_file)
        return True
    
    def process_table_tags(self, tag):
        table_tags = self.container_tag_data.find_all(tag)
        for table_tag in table_tags:
            headers = [th.text.strip() for th in table_tag.find('tr').find_all('th')]
            # print(f"headers:: {headers}")
            table = doc.add_table(rows=1, cols=len(table_tag.find_all('th')))

            #Prepare Header on the table
            row = table.rows[0].cells 
            for index, header in enumerate(headers):
                row[index].text = header

            # Add rows to the table
            for row in table_tag.find_all('tr')[1:]:
                cells = [td.text.strip() for td in row.find_all('td')]
                # print(f"cells: {cells}")
                row = table.add_row().cells 
                for col, cell_text in enumerate(cells):
                    row[col].text = cell_text
            doc.add_paragraph(SPLIT_LINES)

        # doc.save(self.output_file)
        return True

    def process_table_tags_txt_format(self, tag):
        processed_table_tags = ''
        table_tags = self.container_tag_data.find_all(tag)
        for table_tag in table_tags:
            table_data = []
            headers = [th.text.strip() for th in table_tag.find('tr').find_all('th')]
            # table_data.append(headers)
            for row in table_tag.find_all('tr')[1:]:
                row_data = [td.text.strip() for td in row.find_all('td')]
                table_data.append(row_data)
            table_data_in_table_format = tabulate(table_data, headers=headers, tablefmt="grid")
            processed_table_tags += table_data_in_table_format + SPLIT_LINES
            # print(processed_table_tags)
            # print('\n\n')

        doc.add_paragraph(processed_table_tags)
        # doc.save(self.output_file)
        return processed_table_tags
    
    def create_file(self, web_content):
        with open(self.output_file, 'w', encoding='utf-8') as txt_file:
            txt_file.write(web_content)
            txt_file.close()

    def process_website_content(self):
        TAG_FUNCTIONS = {
            'p': self.process_paragraph_tags,
            'img': self.process_image_tags,
            'table': self.process_table_tags
        }

        self.container_tags_data = self.get_container_tag()
        for container_object in self.container_tags_data:
            # print(f"container_object:: {container_object}")
            self.container_tag_data = container_object

            #self.tags => Eg: p, img, table...
            for tag in self.tags:
                tag_exists = TAG_FUNCTIONS.get(tag)
                if tag_exists:
                   tag_exists(tag)
        # OUTPUT_FILES= os.path.abspath(OUTPUT_FILES_DIRECTORY)

        output_files_path = os.path.join(OUTPUT_FILES_DIRECTORY, self.output_file)
        print(f"output_files_path: {output_files_path}")
        doc.save(output_files_path)

# domain_url = "https://www.w3schools.com/html/html_tables.asp"
# container_tag = "body"
# tags = ['p']
# output_filename = 'web'

# scrap_object = ScrapWebiteContent(domain_url, container_tag, tags, output_filename)
# scrap_object.process_website_content()